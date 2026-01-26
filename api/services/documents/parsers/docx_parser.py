# api/services/documents/parsers/docx_parser.py
"""
Parser for Microsoft Word documents (.docx, .doc).
"""

import logging
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import BinaryIO, Optional, Union
import time

from ..base import (
    BaseParser, DocumentContent, DocumentMetadata, 
    FileType, ParserResult, TextChunk
)

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Parser for Microsoft Word documents."""
    
    supported_types = [FileType.DOCX, FileType.DOC]
    
    async def parse(
        self,
        file_path: Union[str, Path],
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
    ) -> ParserResult:
        """Parse a Word document from file path."""
        start_time = time.time()
        file_path = Path(file_path)
        
        try:
            with open(file_path, "rb") as f:
                file_data = f.read()
            
            result = await self._parse_content(
                file_data, 
                file_path.name,
                chunk_size,
                chunk_overlap,
            )
            result.processing_time_ms = (time.time() - start_time) * 1000
            return result
            
        except Exception as e:
            logger.error(f"Error parsing Word document {file_path}: {e}")
            return ParserResult(
                success=False,
                error=str(e),
                processing_time_ms=(time.time() - start_time) * 1000,
            )
    
    async def parse_bytes(
        self,
        file_data: BinaryIO,
        filename: str,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
    ) -> ParserResult:
        """Parse a Word document from bytes."""
        start_time = time.time()
        
        try:
            data = file_data.read()
            result = await self._parse_content(
                data, 
                filename,
                chunk_size,
                chunk_overlap,
            )
            result.processing_time_ms = (time.time() - start_time) * 1000
            return result
            
        except Exception as e:
            logger.error(f"Error parsing Word document bytes: {e}")
            return ParserResult(
                success=False,
                error=str(e),
                processing_time_ms=(time.time() - start_time) * 1000,
            )
    
    async def _parse_content(
        self,
        file_data: bytes,
        filename: str,
        chunk_size: Optional[int],
        chunk_overlap: Optional[int],
    ) -> ParserResult:
        """Parse Word document content."""
        file_type = self.detect_file_type(filename)
        file_hash = self.compute_file_hash(file_data)
        warnings = []
        
        # Write to temp file for processing
        with tempfile.NamedTemporaryFile(
            suffix=f".{file_type.value}", delete=False
        ) as tmp:
            tmp.write(file_data)
            tmp_path = Path(tmp.name)
        
        try:
            # Convert .doc to .docx if needed
            if file_type == FileType.DOC:
                converted_path = await self._convert_doc_to_docx(tmp_path)
                if converted_path:
                    tmp_path = converted_path
                else:
                    warnings.append("Could not convert .doc to .docx, using fallback extraction")
            
            # Extract text using python-docx or pandoc
            full_text, chunks, tables = await self._extract_with_docx(
                tmp_path, chunk_size, chunk_overlap
            )
            
            if not full_text:
                # Fallback to pandoc
                full_text = await self._extract_with_pandoc(tmp_path)
                chunks = self.chunk_text(full_text, "paragraph", chunk_size, chunk_overlap)
            
            # Extract metadata
            metadata = await self._extract_metadata(tmp_path, file_data, filename, file_type, file_hash)
            metadata.word_count = len(full_text.split()) if full_text else 0
            
            content = DocumentContent(
                full_text=full_text,
                chunks=chunks,
                tables=tables,
            )
            
            return ParserResult(
                success=True,
                metadata=metadata,
                content=content,
                warnings=warnings,
            )
            
        finally:
            # Cleanup temp files
            tmp_path.unlink(missing_ok=True)
            if file_type == FileType.DOC:
                docx_path = tmp_path.with_suffix(".docx")
                docx_path.unlink(missing_ok=True)
    
    async def _convert_doc_to_docx(self, doc_path: Path) -> Optional[Path]:
        """Convert .doc to .docx using LibreOffice."""
        try:
            output_dir = doc_path.parent
            result = subprocess.run(
                [
                    "soffice", "--headless", "--convert-to", "docx",
                    "--outdir", str(output_dir), str(doc_path)
                ],
                capture_output=True,
                timeout=60,
            )
            if result.returncode == 0:
                docx_path = doc_path.with_suffix(".docx")
                if docx_path.exists():
                    return docx_path
        except Exception as e:
            logger.warning(f"Failed to convert .doc to .docx: {e}")
        return None
    
    async def _extract_with_docx(
        self,
        file_path: Path,
        chunk_size: Optional[int],
        chunk_overlap: Optional[int],
    ) -> tuple[str, list[TextChunk], list[dict]]:
        """Extract content using python-docx library."""
        try:
            from docx import Document
            from docx.table import Table
        except ImportError:
            logger.warning("python-docx not installed, using fallback")
            return "", [], []
        
        try:
            doc = Document(str(file_path))
            
            paragraphs = []
            chunks = []
            tables = []
            current_heading = None
            chunk_idx = 0
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Track headings for context
                style_name = getattr(para.style, "name", None)
                if style_name and style_name.startswith("Heading"):
                    current_heading = text
                
                paragraphs.append(text)
                
                # Create chunk for this paragraph
                chunks.append(TextChunk(
                    content=text,
                    source=f"paragraph_{chunk_idx}",
                    chunk_index=chunk_idx,
                    heading=current_heading,
                ))
                chunk_idx += 1
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_data = self._extract_table(table)
                if table_data:
                    tables.append({
                        "index": i,
                        "rows": table_data,
                    })
            
            full_text = "\n\n".join(paragraphs)
            
            # Re-chunk if paragraphs are too small or large
            if chunk_size:
                chunks = self.chunk_text(full_text, "chunk", chunk_size, chunk_overlap)
            
            return full_text, chunks, tables
            
        except Exception as e:
            logger.error(f"Error with python-docx: {e}")
            return "", [], []
    
    def _extract_table(self, table) -> list[list[str]]:
        """Extract table data as list of rows."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        return rows
    
    async def _extract_with_pandoc(self, file_path: Path) -> str:
        """Extract text using pandoc as fallback."""
        try:
            result = subprocess.run(
                ["pandoc", "-t", "plain", str(file_path)],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0:
                return result.stdout.strip()
        except Exception as e:
            logger.warning(f"Pandoc extraction failed: {e}")
        return ""
    
    async def _extract_metadata(
        self,
        file_path: Path,
        file_data: bytes,
        filename: str,
        file_type: FileType,
        file_hash: str,
    ) -> DocumentMetadata:
        """Extract document metadata."""
        metadata = DocumentMetadata(
            filename=filename,
            file_type=file_type,
            file_size=len(file_data),
            file_hash=file_hash,
        )
        
        try:
            from docx import Document
            doc = Document(str(file_path))
            
            core_props = doc.core_properties
            metadata.title = core_props.title
            metadata.author = core_props.author
            metadata.subject = core_props.subject
            metadata.created_at = core_props.created
            metadata.modified_at = core_props.modified
            
        except Exception as e:
            logger.warning(f"Could not extract Word metadata: {e}")
        
        return metadata
